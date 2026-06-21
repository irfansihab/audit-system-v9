"""Daftar Temuan & Rekomendasi (DHP, SDP-K.01) — artefak EKSPOR di garis serah.

Dihasilkan OTOMATIS saat laporan disetujui (lihat create_lhp_review) dari
`_KKP/temuan.json` + `_LHP/rekomendasi.json`. Pure-python (python-docx), tanpa LLM.

Inilah satu-satunya "dokumen komunikasi" yang diproduksi v9 sebagai paket ekspor ke
ranah administrasi (Tahapan 8 / TU): tabel temuan (Kondisi-Kriteria-Sebab-Akibat) +
Rekomendasi, dengan kolom Tindak Lanjut / PIC / Target dikosongkan untuk diisi TU.
"""
from __future__ import annotations

import json
from pathlib import Path


def _load_temuan(folder: Path) -> list[dict]:
    p = folder / "_KKP" / "temuan.json"
    if not p.exists():
        return []
    try:
        d = json.loads(p.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return []
    items = d if isinstance(d, list) else (d.get("temuan") or d.get("items") or [])
    return [t for t in items if isinstance(t, dict)]


def _load_rekomendasi(folder: Path) -> dict:
    p = folder / "_LHP" / "rekomendasi.json"
    if not p.exists():
        return {}
    try:
        d = json.loads(p.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return {}
    return d if isinstance(d, dict) else {}


def _rek_text(val) -> str:
    if isinstance(val, str):
        return val
    if isinstance(val, dict):
        return str(val.get("rekomendasi") or val.get("teks") or val.get("text") or "")
    return ""


def build_daftar_temuan_rekomendasi(folder: Path) -> Path | None:
    """Bangun _LHP/Daftar-Temuan-Rekomendasi.docx. Return path, atau None bila
    tak ada temuan (mis. konsultansi/evaluasi-LKE tanpa temuan KKSA)."""
    from docx import Document
    from docx.shared import Pt

    temuan = _load_temuan(folder)
    if not temuan:
        return None
    rek = _load_rekomendasi(folder)

    doc = Document()
    doc.add_heading("DAFTAR TEMUAN DAN REKOMENDASI", level=0)
    doc.add_paragraph(f"Kode penugasan: {folder.name}")
    note = doc.add_paragraph(
        "Dokumen ekspor untuk administrasi & pemantauan tindak lanjut — dihasilkan "
        "otomatis dari Kertas Kerja (KKP) & rekomendasi yang disetujui. Kolom Tindak "
        "Lanjut / PIC / Target diisi oleh Tata Usaha (Tahapan 8)."
    )
    note.runs[0].italic = True

    cols = ["No", "Sasaran", "Uraian Temuan (K-K-S-A)", "Rekomendasi", "Tindak Lanjut", "PIC", "Target"]
    table = doc.add_table(rows=1, cols=len(cols))
    try:
        table.style = "Light Grid Accent 1"
    except Exception:  # noqa: BLE001
        pass
    for i, c in enumerate(cols):
        cell = table.rows[0].cells[i]
        cell.text = c
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    for idx, t in enumerate(temuan, start=1):
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = str(t.get("sasaran_id") or "")
        # Uraian: judul (kode) + K/K/S/A
        uraian = cells[2]
        judul = str(t.get("judul_temuan") or t.get("judul") or "")
        kode = str(t.get("kode_kondisi") or "")
        h = uraian.paragraphs[0]
        run = h.add_run(judul + (f"  [{kode}]" if kode else ""))
        run.bold = True
        for label, key in (("Kondisi", "kondisi"), ("Kriteria", "kriteria"), ("Sebab", "sebab"), ("Akibat", "akibat")):
            val = t.get(key)
            if val in (None, "", []):
                continue
            par = uraian.add_paragraph()
            lab = par.add_run(f"{label}: ")
            lab.bold = True
            par.add_run(str(val))
        cells[3].text = _rek_text(rek.get(t.get("id_temuan") or t.get("id") or ""))
        # cells[4..6] (Tindak Lanjut/PIC/Target) sengaja kosong → diisi TU

    out = folder / "_LHP" / "Daftar-Temuan-Rekomendasi.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out
