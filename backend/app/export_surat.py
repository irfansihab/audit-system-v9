"""Draft Surat Penyampaian LHP (SDP-K, ranah administrasi) — Tahapan 8 / TU.

Menghasilkan DRAFT surat pengantar penyampaian LHP ke klien. TU melengkapi
nomor agenda & tanggal; penomoran resmi, TTE, distribusi, dan arsip TETAP di
SIMWAS (di luar lingkup produksi v9). Pure-python (python-docx), tanpa LLM.
"""
from __future__ import annotations

import json
from pathlib import Path


def _ctx(folder: Path) -> dict:
    p = folder / "context.md"
    out: dict[str, str] = {}
    if p.exists():
        try:
            for line in p.read_text(encoding="utf-8").splitlines():
                if ":" in line and len(line) < 200:
                    k, _, v = line.partition(":")
                    k = k.strip().lower().lstrip("#-* ").strip()
                    if k and v.strip():
                        out.setdefault(k, v.strip())
        except OSError:
            pass
    return out


def _latest_lhp(folder: Path) -> str | None:
    lhp = folder / "_LHP"
    if not lhp.is_dir():
        return None
    docs = [p for p in lhp.glob("*.docx")
            if not p.name.startswith("~$") and "SUBSTANSI" not in p.name
            and "Surat-Penyampaian" not in p.name and "Daftar-Temuan" not in p.name]
    return (max(docs, key=lambda p: p.stat().st_mtime).name) if docs else None


def build_surat_penyampaian(
    folder: Path,
    *,
    nomor: str = "",
    tanggal: str = "",
    tujuan: str = "",
    perihal: str = "",
    auditi: str = "",
) -> Path:
    """Bangun _ADMIN/Surat-Penyampaian-LHP.docx (draft). TU melengkapi sisanya."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ctx = _ctx(folder)
    lhp_nama = _latest_lhp(folder) or "Laporan Hasil Pengawasan"
    auditi = auditi or ctx.get("auditi") or ctx.get("obyek") or ctx.get("objek") or "[Pimpinan Unit Auditi]"
    perihal = perihal or f"Penyampaian {Path(lhp_nama).stem}"

    doc = Document()
    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("SURAT PENYAMPAIAN LAPORAN HASIL PENGAWASAN (DRAFT)"); r.bold = True
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Inspektorat II — Kementerian Komunikasi dan Digital").italic = True

    meta = doc.add_paragraph()
    meta.add_run(f"Nomor\t: {nomor or '[diisi TU / nomor agenda SIMWAS]'}\n")
    meta.add_run(f"Tanggal\t: {tanggal or '[diisi TU]'}\n")
    meta.add_run("Sifat\t: Terbatas\n")
    meta.add_run(f"Hal\t: {perihal}")

    doc.add_paragraph(f"Yth. {tujuan or auditi}")
    doc.add_paragraph(
        "Sehubungan dengan telah selesainya pelaksanaan penugasan pengawasan, "
        f"bersama ini kami sampaikan {Path(lhp_nama).stem} (terlampir) untuk menjadi "
        "perhatian dan ditindaklanjuti sesuai rekomendasi yang termuat di dalamnya."
    )
    doc.add_paragraph(
        "Tindak lanjut atas rekomendasi mohon disampaikan paling lambat 60 (enam puluh) "
        "hari kalender sejak surat ini diterima, disertai bukti pendukung yang memadai."
    )
    doc.add_paragraph("Atas perhatian dan kerja samanya, kami ucapkan terima kasih.")

    ttd = doc.add_paragraph(); ttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ttd.add_run("Inspektur II,\n\n\n\n[Nama]\n[NIP]")

    lamp = doc.add_paragraph(); lamp.add_run("Lampiran:\n").bold = True
    lamp.add_run(f"1. {lhp_nama}\n2. Daftar Temuan dan Rekomendasi")

    note = doc.add_paragraph()
    note.add_run(
        "Catatan: dokumen ini DRAFT dari sistem. Penomoran resmi, tanda tangan elektronik (TTE), "
        "distribusi, dan pengarsipan dilakukan melalui SIMWAS."
    ).italic = True

    out = folder / "_ADMIN" / "Surat-Penyampaian-LHP.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out
