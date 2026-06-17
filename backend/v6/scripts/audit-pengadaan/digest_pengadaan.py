"""
digest_pengadaan.py — Parser konsolidasi dokumen pengadaan → JSON terstruktur.

Scan folder penugasan, identifikasi dokumen berdasarkan pola nama file, lalu
parse setiap dokumen ke schema yang sesuai. Output menyertakan:
  - KAK (Kerangka Acuan Kerja)
  - HPS (Harga Perkiraan Sendiri)
  - Kontrak/SPK/SSUKSSKK
  - SPPBJ (Surat Penunjukan Penyedia)
  - BAST / BA Rekonsiliasi
  - Dokumen Pembayaran (LS, SPTB)

Usage:
    python digest_pengadaan.py <folder-penugasan> [-o output.json]

Pendekatan: regex + filename pattern matching. Cocok untuk audit-pengadaan
dan reviu-pengadaan (sharing schema).
"""

from __future__ import annotations
import argparse
import json
import re
import sys
from pathlib import Path

try:
    import pdfplumber
except ImportError:
    sys.exit("ERROR: pdfplumber not installed. Run: pip install pdfplumber")

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


# ============================================================
# DOCUMENT TYPE CLASSIFICATION (by filename heuristic)
# ============================================================

FILENAME_PATTERNS = {
    "kak": [r"\bKAK\b", r"Kerangka\s+Acuan", r"TOR\b"],
    "hps": [r"\bHPS\b", r"Harga\s+Perkiraan"],
    "hps_detail": [r"Tabel\s+Penyusun.+HPS", r"Rekap\s+Komponen.+HPS"],
    "identifikasi_pengadaan": [r"Identifikasi\s+Pengadaan"],
    "rfi": [r"\bRFI\b", r"Request\s+For\s+Information"],
    "kontrak": [r"SSUKSSKK", r"Salinan\s+Jasa\s+Lainnya", r"Kontrak", r"\bSPK\b"],
    "sppbj": [r"SPPBJ"],
    "perjanjian_kerahasiaan": [r"Perjanjian\s+Kerahasiaan", r"NDA"],
    "permohonan_jaminan": [r"Permohonan\s+.+Jaminan"],
    "pembayaran_ls": [r"\bLS\b.+(Jan|Feb|Mar|Apr|Mei|Jun|Jul|Agu|Sep|Okt|Nov|Des)",
                      r"SPM.+LS", r"Pembayaran\s+LS"],
    "sptb": [r"SPTB", r"Surat\s+Pernyataan\s+Tanggungjawab\s+Belanja"],
    "ba_rekonsiliasi": [r"BA\s+Rekonsiliasi", r"Berita\s+Acara.+SLA"],
    "laporan_bulanan": [r"Laporan\s+Bulanan"],
}


def classify_file(filename: str) -> str | None:
    """Kembalikan jenis dokumen berdasarkan pola nama file."""
    for doc_type, patterns in FILENAME_PATTERNS.items():
        for pat in patterns:
            if re.search(pat, filename, re.I):
                return doc_type
    return None


# ============================================================
# EXTRACTORS (per file type)
# ============================================================

def _extract_pdf_text(path: Path) -> list[str]:
    pages = []
    try:
        with pdfplumber.open(path) as pdf:
            for p in pdf.pages:
                pages.append(p.extract_text() or "")
    except Exception as e:
        return [f"[ERROR: {e}]"]
    return pages


def _extract_docx_text(path: Path) -> list[str]:
    """Ekstrak teks dari docx; kembalikan single-item list untuk kompatibilitas."""
    if DocxDocument is None:
        return ["[docx not supported — install python-docx]"]
    try:
        doc = DocxDocument(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        # tambah teks dari tabel
        for t in doc.tables:
            for row in t.rows:
                text += "\n" + " | ".join(c.text for c in row.cells)
        return [text]
    except Exception as e:
        return [f"[ERROR: {e}]"]


def _extract_text(path: Path) -> list[str]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_text(path)
    elif suffix in (".docx", ".doc"):
        return _extract_docx_text(path)
    return []


def _rupiah_to_int(s: str) -> int | None:
    if not s:
        return None
    m = re.search(r"([\d\.\,]+)", s)
    if not m:
        return None
    num = m.group(1).replace(".", "").replace(",", "")
    try:
        return int(num)
    except ValueError:
        return None


def _extract_periode(text: str) -> str | None:
    """Ekstrak periode pelaksanaan dari teks dokumen pengadaan.

    Aturan ketat untuk hindari false positive:
      1. Prefer pola dengan keyword konteks: "jangka waktu / selama / durasi /
         periode pelaksanaan / masa kontrak / masa pelaksanaan / waktu pelaksanaan".
      2. Bila tidak ada konteks, terima "(\\d{1,3})\\s*(bulan)" dengan angka 1-60.
         JANGAN terima "X tahun" tanpa konteks (terlalu rawan ketabrak nomor PP/UU
         seperti "PP No. 45 Tahun 2013").
      3. Reject angka >99 (jelas bukan durasi proyek).
      4. Reject jika angka berupa tahun (>=1900) atau berdekatan dengan kata
         "Tahun Anggaran".
    """
    # 1) Pola dengan konteks
    ctx_pattern = (
        r"(?:jangka\s+waktu|selama|durasi|periode\s+pelaksanaan|"
        r"masa\s+(?:kontrak|pelaksanaan)|waktu\s+pelaksanaan)"
        r"[^\n]{0,40}?(\d{1,2})\s*(?:\(\w+\))?\s*(bulan|tahun)"
    )
    for m in re.finditer(ctx_pattern, text, re.I):
        num = int(m.group(1))
        unit = m.group(2).lower()
        if unit == "bulan" and 1 <= num <= 60:
            return f"{num} bulan"
        if unit == "tahun" and 1 <= num <= 5:
            return f"{num} tahun"

    # 2) Fallback: hanya pola "X bulan" (1..60), bukan "X tahun"
    for m in re.finditer(r"(?<!\d)(\d{1,2})\s*(?:\(\w+\))?\s*bulan\b", text, re.I):
        num = int(m.group(1))
        if 1 <= num <= 60:
            # Cek 30-char di kiri: jangan kalau didahului "Pasal", "No.", dsb
            left = text[max(0, m.start()-30):m.start()].lower()
            if re.search(r"pasal|no\.|nomor|tahun\s+anggaran|t\.?a\.?\b", left):
                continue
            return f"{num} bulan"

    return None


def parse_kak(pages: list[str]) -> dict:
    text = "\n".join(pages)
    out = {
        "nomor": None, "tanggal": None, "nama_pekerjaan": None,
        "nilai_hps": None, "periode": None, "sla_disebut": False,
        "sla_value": None, "sla_all_values": [], "migrasi_disebut": False,
        "kapasitas_disebut": None, "halaman": len(pages),
    }
    m = re.search(r"Nomor\s*:?\s*(\S+)", text[:1500])
    if m:
        out["nomor"] = m.group(1)[:80]
    m = re.search(r"(\d{1,2}\s+\w+\s+\d{4})", text[:1500])
    if m:
        out["tanggal"] = m.group(1)
    # nama pekerjaan
    m = re.search(r"(?:Pengadaan|Penyediaan|Jasa)\s+([A-Z][^\n.]{10,200})", text[:2000])
    if m:
        out["nama_pekerjaan"] = m.group(0)[:200]
    # periode — STRICT: harus ada konteks ("jangka waktu/selama/durasi/periode/masa")
    # ATAU "X bulan" murni dengan X<=60. Reject angka >99 atau >1900 (tahun rujukan).
    out["periode"] = _extract_periode(text)
    # SLA - scan SELURUH text, tangkap semua angka SLA berbeda (untuk RP.11)
    if re.search(r"\bSLA\b|uptime|service\s+level", text, re.I):
        out["sla_disebut"] = True
        sm = re.search(r"(?:SLA|uptime|service\s+level).{0,80}?(\d{2,3}(?:[,.]\d+)?)\s*%", text, re.I)
        if sm:
            out["sla_value"] = sm.group(1) + "%"
        # Kumpulkan SEMUA angka SLA % yang muncul di seluruh dokumen
        all_sla = re.findall(r"(?:SLA|uptime|service\s+level)[^\n]{0,100}?(\d{2,3}(?:[,.]\d+)?)\s*%", text, re.I)
        # Normalize titik → koma untuk konsistensi pembanding
        unique_sla = sorted(set(s.replace(".", ",") for s in all_sla))
        out["sla_all_values"] = [f"{v}%" for v in unique_sla]
    # migrasi
    if re.search(r"\bmigrasi\b|migration", text, re.I):
        out["migrasi_disebut"] = True
    # kapasitas
    km = re.search(r"(\d+(?:[,.]\d+)?)\s*(Tbps|Gbps|Mbps|TB|GB)", text, re.I)
    if km:
        out["kapasitas_disebut"] = km.group(0)
    # HPS jika ada
    hm = re.search(r"Rp\s*([\d\.,]+)", text)
    if hm:
        out["nilai_hps"] = _rupiah_to_int(hm.group(1))
    # Kelengkapan 5 elemen justifikasi/dokumen persiapan (deteksi keyword full-text).
    # Dipakai rule kelengkapan justifikasi (reviu/audit pengadaan). Heuristik
    # presence-only — bisa false-negative bila frasa tak lazim → rule berseverity
    # PERINGATAN + minta konfirmasi manual.
    out["elemen_justifikasi"] = {
        "kebutuhan": bool(re.search(
            r"latar\s+belakang|identifikasi\s+kebutuhan|analisis\s+kebutuhan|maksud\s+dan\s+tujuan",
            text, re.I)),
        "spesifikasi_teknis": bool(re.search(
            r"spesifikasi\s+teknis|persyaratan\s+teknis|spesifikasi\s+barang|spek\s+teknis|"
            r"spesifikasi\s+fungsi|kebutuhan\s+fungsional",
            text, re.I)),
        "metode_pengadaan": bool(re.search(
            r"metode\s+(?:pengadaan|pemilihan)|cara\s+pengadaan|tender\b|seleksi\b|"
            r"penunjukan\s+langsung|pengadaan\s+langsung|e-?purchasing|e-?katalog|e-?katalogue",
            text, re.I)),
        # waktu: reuse hasil _extract_periode atau frasa jadwal/waktu penyelesaian
        "waktu_penyelesaian": bool(out["periode"]) or bool(re.search(
            r"jangka\s+waktu|waktu\s+pelaksanaan|jadwal\s+pelaksanaan|masa\s+pelaksanaan|"
            r"waktu\s+penyelesaian|jadwal\s+kegiatan",
            text, re.I)),
        "output": bool(re.search(
            r"keluaran|output|deliverable|hasil\s+yang\s+diharapkan|hasil\s+pekerjaan|"
            r"produk\s+yang\s+dihasilkan|sasaran\s+keluaran",
            text, re.I)),
    }
    return out


def parse_hps(pages: list[str]) -> dict:
    text = "\n".join(pages)
    out = {
        "nomor": None, "tanggal": None, "total": None,
        "komponen_count": 0, "sla_disebut": False, "sla_value": None,
        "sla_all_values": [],
        "periode": None, "migrasi_disebut": False,
        "ada_dokumen_pembentuk_harga": False, "halaman": len(pages),
    }
    m = re.search(r"Nomor\s*:?\s*(\S+)", text[:1500])
    if m:
        out["nomor"] = m.group(1)[:80]
    # total HPS — biasanya ada "Total" atau "Jumlah" besar
    tm = re.search(r"(?:Total|Jumlah|Grand\s+Total)\s+(?:Rp\s*)?([\d\.,]{10,})", text)
    if tm:
        out["total"] = _rupiah_to_int(tm.group(1))
    # komponen count — baris dengan "Rp" di akhir
    rupiah_lines = re.findall(r"Rp\s*[\d\.,]{6,}", text)
    out["komponen_count"] = len(rupiah_lines)
    # SLA
    if re.search(r"\bSLA\b|uptime|service\s+level", text, re.I):
        out["sla_disebut"] = True
        sm = re.search(r"(?:SLA|uptime|service\s+level).{0,80}?(\d{2,3}(?:[,.]\d+)?)\s*%", text, re.I)
        if sm:
            out["sla_value"] = sm.group(1) + "%"
        all_sla = re.findall(r"(?:SLA|uptime|service\s+level)[^\n]{0,100}?(\d{2,3}(?:[,.]\d+)?)\s*%", text, re.I)
        unique_sla = sorted(set(s.replace(".", ",") for s in all_sla))
        out["sla_all_values"] = [f"{v}%" for v in unique_sla]
    # periode — STRICT (lihat _extract_periode)
    out["periode"] = _extract_periode(text)
    # migrasi
    if re.search(r"\bmigrasi\b|migration", text, re.I):
        out["migrasi_disebut"] = True
    # dokumen pembentuk harga (quotation vendor, market research)
    if re.search(r"penawaran\s+vendor|quotation|market\s+research|RFI", text, re.I):
        out["ada_dokumen_pembentuk_harga"] = True
    return out


def parse_kontrak(pages: list[str]) -> dict:
    text = "\n".join(pages)
    out = {
        "nomor": None, "tanggal": None, "nilai_kontrak": None,
        "penyedia": None, "periode_mulai": None, "periode_selesai": None,
        "sla_clause": False, "sla_value": None, "jaminan_pelaksanaan": None,
        "metode_pembayaran": None, "halaman": len(pages),
    }
    m = re.search(r"Nomor\s*:?\s*(\S+)", text[:1500])
    if m:
        out["nomor"] = m.group(1)[:80]
    # nilai kontrak
    nm = re.search(r"(?:Nilai\s+Kontrak|Total\s+Kontrak|sebesar|senilai)[^\n]{0,50}?Rp\s*([\d\.\,]+)", text, re.I)
    if nm:
        out["nilai_kontrak"] = _rupiah_to_int(nm.group(1))
    # penyedia
    pm = re.search(r"(?:Penyedia|Kontraktor)\s*:?\s*([A-Z][A-Za-z\s\.\,&]{5,80})", text)
    if pm:
        out["penyedia"] = pm.group(1).strip()[:80]
    # SLA
    if re.search(r"\bSLA\b|Service\s+Level\s+Agreement", text, re.I):
        out["sla_clause"] = True
        sm = re.search(r"SLA.{0,80}?(\d{2,3}(?:[,.]\d+)?)\s*%", text, re.I)
        if sm:
            out["sla_value"] = sm.group(1) + "%"
    # jaminan pelaksanaan (biasanya 5%)
    jm = re.search(r"[Jj]aminan\s+[Pp]elaksanaan.{0,100}?(\d+)\s*%", text)
    if jm:
        out["jaminan_pelaksanaan"] = f"{jm.group(1)}%"
    # metode pembayaran
    if re.search(r"\bLS\b|Langsung|Lump.?sum", text):
        out["metode_pembayaran"] = "LS/Lump-sum"
    elif re.search(r"[Bb]ulanan|termin", text):
        out["metode_pembayaran"] = "Termin/Bulanan"
    return out


def parse_bast(pages: list[str]) -> dict:
    text = "\n".join(pages)
    out = {
        "nomor": None, "tanggal": None, "periode": None,
        "nilai_diterima": None, "sla_reported": None,
        "sla_passed": None, "halaman": len(pages),
    }
    m = re.search(r"Nomor\s*:?\s*(\S+)", text[:1500])
    if m:
        out["nomor"] = m.group(1)[:80]
    # periode reconciled
    m = re.search(r"(Januari|Februari|Maret|April|Mei|Juni|Juli|Agustus|September|Oktober|November|Desember)\s*(\d{4})?", text, re.I)
    if m:
        out["periode"] = m.group(0)
    # SLA reported
    sm = re.search(r"SLA.{0,80}?(\d{2,3}(?:[,.]\d+)?)\s*%", text, re.I)
    if sm:
        out["sla_reported"] = sm.group(1) + "%"
    return out


def parse_pembayaran(pages: list[str]) -> dict:
    text = "\n".join(pages)
    out = {
        "nomor": None, "tanggal": None, "nilai": None,
        "periode": None, "bukti_pendukung_lengkap": False,
        "halaman": len(pages),
    }
    m = re.search(r"Nomor\s*:?\s*(\S+)", text[:1500])
    if m:
        out["nomor"] = m.group(1)[:80]
    nm = re.search(r"(?:sebesar|senilai)[^\n]{0,50}?Rp\s*([\d\.\,]+)", text, re.I)
    if nm:
        out["nilai"] = _rupiah_to_int(nm.group(1))
    # bukti pendukung
    if re.search(r"(?:BAST|Berita\s+Acara|BA\s+Rekonsiliasi|SPTB|Invoice|Kwitansi)", text, re.I):
        out["bukti_pendukung_lengkap"] = True
    return out


PARSERS = {
    "kak": parse_kak,
    "hps": parse_hps,
    "kontrak": parse_kontrak,
    "ba_rekonsiliasi": parse_bast,
    "pembayaran_ls": parse_pembayaran,
    "sptb": parse_pembayaran,
}


# ============================================================
# RFI PARSER (untuk multi-source HPS validation)
# ============================================================

_RE_REFUSAL = re.compile(
    r"(belum\s+dapat\s+(?:berpartisipasi|ber\s*partisipasi|ikut|mengikuti)|"
    r"tidak\s+dapat\s+(?:berpartisipasi|memberikan|menyampaikan|ikut)|"
    r"menolak|"
    r"decline|not\s+interested|cannot\s+participate|unable\s+to\s+participate|"
    r"belum\s+sanggup|tidak\s+bersedia)",
    re.I,
)
_RE_PRICE = re.compile(r"(Rp\s*[\d.,]+|IDR\s*[\d.,]+|USD\s*[\d.,]+|harga\s+(?:total|estimasi|perkiraan))", re.I)


def parse_rfi(pages: list[str]) -> dict:
    """Parse RFI: deteksi apakah vendor memberikan harga atau menolak partisipasi."""
    text = "\n".join(pages) if pages else ""
    refusal = bool(_RE_REFUSAL.search(text))
    has_price = bool(_RE_PRICE.search(text))
    # Vendor: cari nama vendor di header atau signature
    vendor_match = re.search(r"(?:^|\n)\s*(?:PT\.?\s+|CV\.?\s+|VP\s+|Hormat\s+kami,?\s*\n+\s*)([A-Z][A-Z\s&.,-]{4,60})", text)
    vendor = vendor_match.group(1).strip() if vendor_match else None
    return {
        "vendor_terdeteksi": vendor,
        "memberikan_harga": has_price and not refusal,
        "menolak_partisipasi": refusal,
        "punya_indikasi_harga": has_price,
        "halaman": len(pages) if pages else 0,
    }


PARSERS["rfi"] = parse_rfi


# ============================================================
# FOLDER SCANNER
# ============================================================

def scan_folder(folder: Path) -> dict:
    """Scan folder, klasifikasi + parse setiap dokumen."""
    out = {
        "metadata": {
            "folder_source": str(folder),
            "parser_version": "v0.2",
        },
        "dokumen": {},
        "unclassified_files": [],
        "missing_types": [],
    }

    # walk
    all_files = list(folder.rglob("*"))
    for f in all_files:
        if not f.is_file():
            continue
        if f.suffix.lower() not in (".pdf", ".docx", ".doc"):
            continue
        if f.name == "desktop.ini" or f.name.startswith("~"):
            continue

        doc_type = classify_file(f.name)
        if not doc_type:
            out["unclassified_files"].append(str(f.relative_to(folder)))
            continue

        parser = PARSERS.get(doc_type)
        entry = {
            "filename": f.name,
            "path": str(f.relative_to(folder)),
            "jenis_dokumen": doc_type,
            "parsed": None,
        }
        if parser:
            try:
                pages = _extract_text(f)
                entry["parsed"] = parser(pages)
                entry["parsed"]["_raw_first_chars"] = ("\n".join(pages))[:2500] if pages else ""
            except Exception as e:
                entry["parsed"] = {"_error": str(e)}

        out["dokumen"].setdefault(doc_type, []).append(entry)

    # Cek missing types yang penting
    penting = ["kak", "hps", "kontrak"]
    for t in penting:
        if t not in out["dokumen"]:
            out["missing_types"].append(t)

    return out


def _self_check_ast() -> None:
    """Preflight: pastikan script ini sendiri syntactically valid sebelum run."""
    import ast
    try:
        ast.parse(open(__file__, "r", encoding="utf-8").read())
    except SyntaxError as e:
        print(f"Self-check AST gagal di {__file__}: {e}", file=sys.stderr)
        print("   File mungkin korup. Lihat backup atau git restore.", file=sys.stderr)
        sys.exit(2)


def main(argv=None):
    _self_check_ast()
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("input_dir", help="Folder berisi PDF/dokumen pengadaan")
    ap.add_argument("-o", "--output", default="pengadaan-digest.json")
    args = ap.parse_args(argv)

    in_dir = Path(args.input_dir)
    if not in_dir.exists():
        print(f"Input dir tidak ditemukan: {in_dir}", file=sys.stderr)
        return 1

    digest = scan_folder(in_dir)
    Path(args.output).write_text(json.dumps(digest, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"OK: {args.output}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    sys.exit(main())
